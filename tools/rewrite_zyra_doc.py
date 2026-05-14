from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "zyra Project Doc.docx"
REPORT_PATH = ROOT / "Zyra_Project_Report.md"
BACKUP_PATH = ROOT / "zyra Project Doc.backup.docx"

ABSTRACT_TEXT = (
    "Zyra is an AI-powered recruitment, screening, and interview management system "
    "developed to automate the early stages of hiring through a structured, web-based "
    "workflow. The platform allows HR teams to publish roles, collect candidate "
    "applications, perform resume screening, conduct AI-generated MCQ assessments, "
    "enable technical coding rounds for suitable candidates, and run AI avatar-based "
    "interviews with proctoring support. The system combines Flask, MongoDB, resume "
    "analysis, Groq and optional Ollama-based question generation, Cloudinary and GridFS "
    "storage support, and role-based HR reporting to reduce manual effort and improve "
    "consistency in candidate evaluation. Zyra also emphasizes fairness and security by "
    "focusing on role-relevant evidence, restricting final recommendations to HR users, "
    "tracking proctoring violations, and preserving candidate records for later review. "
    "The result is a practical hiring platform that improves screening efficiency, "
    "supports better decision-making, and demonstrates how AI can be integrated into "
    "real-world recruitment workflows."
)

KEYWORDS_TEXT = (
    "Keywords: AI Recruitment, Resume Screening, MCQ Assessment, Coding Round, "
    "Avatar Interview, Proctoring, HR Dashboard, Flask, MongoDB."
)


def normalize_ascii(text: str) -> str:
    replacements = {
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
        "\u00a0": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def clean_inline(text: str) -> str:
    text = normalize_ascii(text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.strip()


def set_run_font(run, size_pt: int, *, bold: bool = False, italic: bool = False) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def write_paragraph(paragraph, text: str, kind: str) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if kind == "chapter" else WD_ALIGN_PARAGRAPH.LEFT

    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)

    if kind == "chapter":
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(10)
    elif kind in {"heading", "subheading"}:
        fmt.space_before = Pt(8)
        fmt.space_after = Pt(6)
    elif kind == "code":
        fmt.left_indent = Inches(0.35)
        fmt.space_after = Pt(3)
    elif kind in {"bullet", "number"}:
        fmt.left_indent = Inches(0.3)
        fmt.first_line_indent = Inches(0)

    run = paragraph.add_run(clean_inline(text))
    if kind == "chapter":
        set_run_font(run, 16, bold=True)
    elif kind in {"heading", "subheading"}:
        set_run_font(run, 14, bold=True)
    else:
        set_run_font(run, 12)


def style_existing_paragraph(paragraph, kind: str, text: str) -> None:
    write_paragraph(paragraph, text, kind)


def find_paragraph_index(doc: Document, target_text: str) -> int:
    wanted = clean_inline(target_text).lower()
    for idx, para in enumerate(doc.paragraphs):
        if clean_inline(para.text).lower() == wanted:
            return idx
    raise ValueError(f"Paragraph not found: {target_text}")


def find_first_existing_paragraph_index(doc: Document, *targets: str) -> int:
    for target in targets:
        try:
            return find_paragraph_index(doc, target)
        except ValueError:
            continue
    raise ValueError(f"None of the paragraph targets were found: {targets}")


def ensure_abstract_paragraphs(doc: Document) -> tuple:
    abstract_idx = find_paragraph_index(doc, "ABSTRACT")
    toc_idx = find_paragraph_index(doc, "TABLE OF CONTENTS")

    non_empty = [p for p in doc.paragraphs[abstract_idx + 1 : toc_idx] if p.text.strip()]
    toc_para = doc.paragraphs[toc_idx]

    while len(non_empty) < 2:
        inserted = toc_para.insert_paragraph_before("")
        non_empty.append(inserted)

    first, second = non_empty[0], non_empty[1]
    for extra in non_empty[2:]:
        clear_paragraph(extra)

    style_existing_paragraph(doc.paragraphs[abstract_idx], "chapter", "ABSTRACT")
    style_existing_paragraph(first, "body", ABSTRACT_TEXT)
    style_existing_paragraph(second, "body", KEYWORDS_TEXT)
    return first, second


def remove_body_from_chapter_one(doc: Document) -> None:
    start_idx = find_first_existing_paragraph_index(
        doc,
        "CHAPTER 1 INTRODUCTION",
        "CHAPTER 1: INTRODUCTION",
    )
    start_element = doc.paragraphs[start_idx]._element
    body = doc._element.body

    removing = False
    for child in list(body):
        if child is start_element:
            removing = True
        if removing and child.tag != qn("w:sectPr"):
            body.remove(child)


def parse_markdown_blocks(lines: list[str]) -> list[tuple[str, object]]:
    blocks: list[tuple[str, object]] = []
    i = 0
    in_code = False
    code_lines: list[str] = []

    def is_special(value: str) -> bool:
        stripped = value.strip()
        return (
            not stripped
            or stripped.startswith("# ")
            or stripped.startswith("## ")
            or stripped.startswith("### ")
            or stripped.startswith("- ")
            or re.match(r"^\d+\.\s+", stripped) is not None
            or stripped.startswith("|")
            or stripped == "---"
            or stripped.startswith("```")
        )

    while i < len(lines):
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                blocks.append(("code", code_lines[:]))
                code_lines.clear()
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(normalize_ascii(raw.rstrip()))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            blocks.append(("h1", clean_inline(stripped[2:])))
            i += 1
            continue

        if stripped.startswith("## "):
            blocks.append(("h2", clean_inline(stripped[3:])))
            i += 1
            continue

        if stripped.startswith("### "):
            blocks.append(("h3", clean_inline(stripped[4:])))
            i += 1
            continue

        if stripped.startswith("|"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if not re.fullmatch(r"\|?[\s:-]+\|[\s|:-]*", row_line):
                    row = [clean_inline(cell) for cell in row_line.strip("|").split("|")]
                    rows.append(row)
                i += 1
            if rows:
                blocks.append(("table", rows))
            continue

        if stripped.startswith("- "):
            items: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(clean_inline(lines[i].strip()[2:]))
                i += 1
            blocks.append(("ul", items))
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items: list[str] = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(clean_inline(re.sub(r"^\d+\.\s+", "", lines[i].strip())))
                i += 1
            blocks.append(("ol", items))
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines) and not is_special(lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        blocks.append(("p", clean_inline(" ".join(para_lines))))

    return blocks


def format_cell_paragraph(paragraph, text: str, *, header: bool = False) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(3)
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, 12, bold=header)


def append_blocks(doc: Document, blocks: list[tuple[str, object]]) -> None:
    chapter_started = False
    for kind, payload in blocks:
        if kind == "h1":
            doc.add_page_break()
            p = doc.add_paragraph()
            write_paragraph(p, str(payload), "chapter")
            chapter_started = True
            continue

        if kind == "h2":
            p = doc.add_paragraph()
            write_paragraph(p, str(payload), "heading")
            continue

        if kind == "h3":
            p = doc.add_paragraph()
            write_paragraph(p, str(payload), "subheading")
            continue

        if kind == "p":
            p = doc.add_paragraph()
            write_paragraph(p, str(payload), "body")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue

        if kind == "ul":
            for item in payload:
                p = doc.add_paragraph()
                write_paragraph(p, f"- {item}", "bullet")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue

        if kind == "ol":
            for idx, item in enumerate(payload, start=1):
                p = doc.add_paragraph()
                write_paragraph(p, f"{idx}. {item}", "number")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue

        if kind == "code":
            for line in payload:
                p = doc.add_paragraph()
                write_paragraph(p, line if line else " ", "code")
            continue

        if kind == "table":
            rows = payload
            col_count = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            for r_idx, row in enumerate(rows):
                for c_idx in range(col_count):
                    value = row[c_idx] if c_idx < len(row) else ""
                    cell = table.cell(r_idx, c_idx)
                    format_cell_paragraph(cell.paragraphs[0], value, header=(r_idx == 0))
            doc.add_paragraph()
            continue

    if not chapter_started:
        raise ValueError("No chapter content was appended from the markdown report.")


def extract_report_chapters() -> list[tuple[str, object]]:
    lines = REPORT_PATH.read_text(encoding="utf-8").splitlines()
    start = None
    for idx, line in enumerate(lines):
        if line.strip().startswith("# CHAPTER 1"):
            start = idx
            break
    if start is None:
        raise ValueError("Could not find Chapter 1 in Zyra_Project_Report.md")
    return parse_markdown_blocks(lines[start:])


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(f"Missing document: {DOCX_PATH}")
    if not REPORT_PATH.exists():
        raise FileNotFoundError(f"Missing report: {REPORT_PATH}")

    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(str(DOCX_PATH))
    ensure_abstract_paragraphs(doc)
    remove_body_from_chapter_one(doc)
    blocks = extract_report_chapters()
    append_blocks(doc, blocks)
    doc.save(str(DOCX_PATH))
    print(f"Updated document: {DOCX_PATH}")
    print(f"Backup kept at: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
